"""Helpers for extracting plain text from various document types."""
from __future__ import annotations

import html
from pathlib import Path

from pypdf import PdfReader

from doc_to_md.utils.validation import (
    is_likely_corrupted_docx,
    is_likely_corrupted_pdf,
    validate_file,
)

try:  # Optional dependency for DOCX parsing
    from docx import Document
except ImportError:  # pragma: no cover - handled in runtime guard
    Document = None  # type: ignore[assignment]

try:  # Optional dependency for image OCR
    from PIL import Image
except ImportError:  # pragma: no cover - handled in runtime guard
    Image = None  # type: ignore[assignment]

try:  # Optional dependency for image OCR
    import pytesseract
except ImportError:  # pragma: no cover - handled in runtime guard
    pytesseract = None  # type: ignore[assignment]


def extract_text(path: Path) -> str:
    """
    Extract text from various document formats with validation and error handling.
    
    Args:
        path: Path to the document file
        
    Returns:
        Extracted text content
        
    Raises:
        FileValidationError: If file validation fails
        RuntimeError: If extraction fails due to missing dependencies
    """
    # Validate file before attempting extraction
    validate_file(path)
    
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix == ".docx":
            return _extract_docx(path)
        if suffix in {".png", ".jpg", ".jpeg"}:
            return _extract_image(path)
        return _extract_text_file(path)
    except Exception as exc:
        # Provide helpful error message
        raise RuntimeError(
            f"Failed to extract text from {path.name}: {exc}"
        ) from exc


def _extract_pdf(path: Path) -> str:
    """
    Extract text from PDF with corruption detection and error handling.
    
    Args:
        path: Path to PDF file
        
    Returns:
        Extracted text from all pages
    """
    # Check for obvious corruption
    if is_likely_corrupted_pdf(path):
        return "[PDF file appears to be corrupted or invalid]"
    
    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001
        return f"[PDF reading failed: {exc}]"
    
    if len(reader.pages) == 0:
        return "[PDF contains no pages]"
    
    parts: list[str] = []
    failed_pages = 0
    
    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 - best effort
            page_text = f"[Page {index} extraction failed: {exc}]"
            failed_pages += 1
        
        if page_text.strip():
            # Escape special Markdown characters in extracted text
            parts.append(_escape_markdown_special_chars(page_text.strip()))
    
    if not parts:
        return "[No textual content could be extracted from PDF]"
    
    result = "\n\n".join(parts)
    
    # Add warning if many pages failed
    if failed_pages > 0:
        warning = f"\n\n_Note: {failed_pages} page(s) failed to extract._"
        result = warning + "\n\n" + result
    
    return result


def _escape_markdown_special_chars(text: str) -> str:
    """
    Escape special characters that have meaning in Markdown.
    
    Args:
        text: Raw text to escape
        
    Returns:
        Text with special characters escaped
    """
    # Escape HTML entities first
    text = html.escape(text, quote=False)
    
    # Don't escape characters that are commonly part of natural text
    # Only escape when they appear in problematic contexts
    # This is a balanced approach - too much escaping makes text unreadable
    
    return text


def _extract_text_file(path: Path) -> str:
    """
    Extract text from plain text files with proper error handling.
    
    Args:
        path: Path to text file
        
    Returns:
        File contents
    """
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        return f"[Text file reading failed: {exc}]"


def _extract_docx(path: Path) -> str:
    """
    Extract text from DOCX with corruption detection and improved table handling.
    
    Args:
        path: Path to DOCX file
        
    Returns:
        Extracted text with proper Markdown table formatting
    """
    if Document is None:  # pragma: no cover - exercised only when dependency missing
        raise RuntimeError(
            "DOCX extraction requires the 'python-docx' package. "
            "Install project dependencies or run `pip install python-docx`."
        )
    
    # Check for corruption
    if is_likely_corrupted_docx(path):
        return "[DOCX file appears to be corrupted or invalid]"
    
    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001 - propagate friendly message
        return f"[DOCX extraction failed: {exc}]"

    blocks: list[str] = []
    
    # Extract paragraphs with heading detection
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        # Detect heading styles and convert to Markdown
        if paragraph.style.name.startswith("Heading"):
            try:
                level = int(paragraph.style.name.replace("Heading ", ""))
                text = f"{'#' * level} {text}"
            except ValueError:
                pass  # Not a numbered heading style
        
        blocks.append(text)

    # Extract tables with proper Markdown table formatting
    for table in document.tables:
        table_md = _format_table_as_markdown(table)
        if table_md:
            blocks.append(table_md)

    if not blocks:
        return "[No textual content found in DOCX]"
    
    return "\n\n".join(blocks)


def _format_table_as_markdown(table) -> str:
    """
    Format a DOCX table as a Markdown table.
    
    Args:
        table: python-docx Table object
        
    Returns:
        Markdown-formatted table string
    """
    if not table.rows:
        return ""
    
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()).strip() for cell in row.cells]
        rows.append(cells)
    
    if not rows:
        return ""
    
    # Build Markdown table
    md_lines = []
    
    # Header row
    header = rows[0]
    md_lines.append("| " + " | ".join(header) + " |")
    
    # Separator row
    md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    
    # Data rows
    for row in rows[1:]:
        # Pad row to match header length
        while len(row) < len(header):
            row.append("")
        md_lines.append("| " + " | ".join(row[:len(header)]) + " |")
    
    return "\n".join(md_lines)


def _extract_image(path: Path) -> str:
    """
    Extract text from images using OCR with proper error handling.
    
    Args:
        path: Path to image file
        
    Returns:
        OCR-extracted text
    """
    if Image is None or pytesseract is None:  # pragma: no cover - dependency guard
        raise RuntimeError(
            "Image OCR requires Pillow and pytesseract. "
            "Install project dependencies or run `pip install pillow pytesseract`."
        )

    try:
        with Image.open(path) as image:
            # Validate image
            if image.width * image.height > 100_000_000:  # 100 megapixels
                return "[Image too large for OCR (>100 megapixels)]"
            
            grayscale = image.convert("L")  # grayscale improves OCR signal
            text = pytesseract.image_to_string(grayscale)
    except Exception as exc:  # noqa: BLE001 - best effort
        return f"[Image OCR failed: {exc}]"

    return text.strip() or "[No textual content detected in image]"
