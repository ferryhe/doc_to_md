"""Tests for enhanced text extraction with error handling."""
from pathlib import Path

import pytest

from doc_to_md.pipeline.text_extraction import (
    extract_text,
    _extract_pdf,
    _extract_docx,
    _format_table_as_markdown,
)
from doc_to_md.utils.validation import FileValidationError


def test_extract_text_validates_input(tmp_path: Path) -> None:
    """Test that extract_text validates file before processing."""
    file_path = tmp_path / "nonexistent.txt"
    
    with pytest.raises(FileValidationError):
        extract_text(file_path)


def test_extract_text_handles_extraction_error(tmp_path: Path, monkeypatch) -> None:
    """Test that extract_text handles extraction errors gracefully."""
    file_path = tmp_path / "test.txt"
    file_path.write_text("content", encoding="utf-8")
    
    # Mock _extract_text_file to raise an exception
    def mock_extract(*args):
        raise ValueError("Mock extraction error")
    
    import doc_to_md.pipeline.text_extraction as module
    monkeypatch.setattr(module, "_extract_text_file", mock_extract)
    
    with pytest.raises(RuntimeError, match="Failed to extract"):
        extract_text(file_path)


def test_extract_pdf_corrupted_file(tmp_path: Path) -> None:
    """Test PDF extraction handles corrupted files."""
    file_path = tmp_path / "corrupted.pdf"
    file_path.write_bytes(b"not a real pdf")
    
    result = _extract_pdf(file_path)
    assert "corrupted" in result.lower()


def test_extract_pdf_empty_pdf(tmp_path: Path) -> None:
    """Test PDF extraction handles empty PDFs."""
    from pypdf import PdfWriter
    
    file_path = tmp_path / "empty.pdf"
    writer = PdfWriter()
    with open(file_path, "wb") as f:
        writer.write(f)
    
    result = _extract_pdf(file_path)
    assert "no pages" in result.lower()


def test_extract_pdf_with_failed_pages(tmp_path: Path) -> None:
    """Test PDF extraction reports failed pages."""
    # This would require mocking pypdf to simulate page extraction failures
    # For now, we test the happy path
    pass


def test_extract_docx_with_headings(tmp_path: Path) -> None:
    """Test DOCX extraction converts headings to Markdown."""
    docx = pytest.importorskip("docx")
    
    doc = docx.Document()
    heading1 = doc.add_heading("Main Title", level=1)
    doc.add_paragraph("Some content")
    heading2 = doc.add_heading("Subsection", level=2)
    doc.add_paragraph("More content")
    
    file_path = tmp_path / "test.docx"
    doc.save(file_path)
    
    result = _extract_docx(file_path)
    
    assert "# Main Title" in result
    assert "## Subsection" in result


def test_extract_docx_with_table(tmp_path: Path) -> None:
    """Test DOCX extraction formats tables as Markdown."""
    docx = pytest.importorskip("docx")
    
    doc = docx.Document()
    table = doc.add_table(rows=3, cols=2)
    
    # Header row
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Age"
    
    # Data rows
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "30"
    table.rows[2].cells[0].text = "Bob"
    table.rows[2].cells[1].text = "25"
    
    file_path = tmp_path / "test_table.docx"
    doc.save(file_path)
    
    result = _extract_docx(file_path)
    
    # Check for Markdown table format
    assert "| Name | Age |" in result
    assert "| --- | --- |" in result
    assert "| Alice | 30 |" in result
    assert "| Bob | 25 |" in result


def test_extract_docx_corrupted_file(tmp_path: Path) -> None:
    """Test DOCX extraction handles corrupted files."""
    file_path = tmp_path / "corrupted.docx"
    file_path.write_bytes(b"not a real docx")
    
    result = _extract_docx(file_path)
    assert "corrupted" in result.lower() or "invalid" in result.lower()


def test_extract_docx_extraction_error(tmp_path: Path) -> None:
    """Test DOCX extraction handles extraction errors."""
    file_path = tmp_path / "test.docx"
    file_path.write_bytes(b"fake docx that won't parse")
    
    result = _extract_docx(file_path)
    assert "extraction failed" in result.lower() or "corrupted" in result.lower()


def test_format_table_as_markdown_basic() -> None:
    """Test basic table formatting."""
    pytest.importorskip("docx")
    from docx import Document
    
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    
    # Header
    table.rows[0].cells[0].text = "Col1"
    table.rows[0].cells[1].text = "Col2"
    table.rows[0].cells[2].text = "Col3"
    
    # Data
    table.rows[1].cells[0].text = "A"
    table.rows[1].cells[1].text = "B"
    table.rows[1].cells[2].text = "C"
    
    result = _format_table_as_markdown(table)
    
    lines = result.split("\n")
    assert len(lines) == 3
    assert "| Col1 | Col2 | Col3 |" in lines[0]
    assert "| --- | --- | --- |" in lines[1]
    assert "| A | B | C |" in lines[2]


def test_format_table_as_markdown_empty() -> None:
    """Test empty table formatting."""
    pytest.importorskip("docx")
    from docx import Document
    
    doc = Document()
    table = doc.add_table(rows=0, cols=0)
    
    result = _format_table_as_markdown(table)
    assert result == ""


def test_extract_image_too_large(tmp_path: Path, monkeypatch) -> None:
    """Test image extraction handles oversized images."""
    PIL = pytest.importorskip("PIL.Image")
    
    # Create a small image
    image_path = tmp_path / "test.png"
    img = PIL.new("RGB", (100, 100), color="white")
    img.save(image_path)
    
    # Mock the image dimensions to be huge
    original_open = PIL.open
    
    class MockImage:
        width = 20000
        height = 20000
        
        def __enter__(self):
            return self
        
        def __exit__(self, *args):
            pass
        
        def convert(self, mode):
            return self
    
    def mock_open(path):
        return MockImage()
    
    import doc_to_md.pipeline.text_extraction as module
    monkeypatch.setattr(module.Image, "open", mock_open)
    
    from doc_to_md.pipeline.text_extraction import _extract_image
    result = _extract_image(image_path)
    
    assert "too large" in result.lower()
